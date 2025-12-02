import fitz  # PyMuPDF
from PIL import Image
import io
from docx import Document
import pandas as pd

def pdf_to_images(file_bytes):
    """แปลง PDF เป็น List ของรูปภาพ"""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    images = []
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        pix = page.get_pixmap(dpi=150)
        img = Image.open(io.BytesIO(pix.tobytes()))
        images.append(img)
    return images

def create_word(text_list):
    """สร้างไฟล์ Word จาก List ข้อความ"""
    doc = Document()
    for i, text in enumerate(text_list):
        doc.add_heading(f'Page {i+1}', level=1)
        doc.add_paragraph(str(text)) # กันเหนียวแปลงเป็น str
        doc.add_page_break()
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_excel(data_dict):
    """สร้างไฟล์ Excel จาก Dict {page_idx: csv_string}"""
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
        has_data = False
        for sheet_name, csv_data in data_dict.items():
            try:
                df = pd.read_csv(io.StringIO(csv_data))
                df.to_excel(writer, sheet_name=sheet_name, index=False)
                has_data = True
            except:
                pass
        if not has_data:
            pd.DataFrame({"Status": ["No Data"]}).to_excel(writer, sheet_name="Info")
    buffer.seek(0)
    return buffer
